"""
src/parser.py
Handles extraction of text from PDF, DOCX, and plain text inputs.
"""

import io
import os
from pathlib import Path
from typing import Union


def extract_text_from_pdf(source: Union[str, bytes, Path]) -> str:
    """Extract text from a PDF file path or bytes."""
    import pdfplumber

    if isinstance(source, (str, Path)):
        with pdfplumber.open(source) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    elif isinstance(source, bytes):
        with pdfplumber.open(io.BytesIO(source)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    else:
        raise ValueError("source must be a file path or bytes")

    return "\n".join(pages).strip()


def extract_text_from_docx(source: Union[str, bytes, Path]) -> str:
    """Extract text from a DOCX file path or bytes."""
    from docx import Document

    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs).strip()


def extract_text(source: Union[str, bytes, Path], filename: str = "") -> str:
    """
    Auto-detect file type and extract text.
    Falls back to treating source as plain text if it's already a string.
    """
    if isinstance(source, str) and not os.path.exists(source):
        # Already plain text
        return source.strip()

    ext = Path(filename or str(source)).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(source)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(source)
    elif ext in (".txt", ".md"):
        if isinstance(source, bytes):
            return source.decode("utf-8", errors="ignore").strip()
        return Path(source).read_text(encoding="utf-8", errors="ignore").strip()
    else:
        # Try PDF first, then DOCX
        try:
            return extract_text_from_pdf(source)
        except Exception:
            try:
                return extract_text_from_docx(source)
            except Exception:
                raise ValueError(f"Unsupported file type: {ext or 'unknown'}")
